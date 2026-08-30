from pathlib import Path
from docx import Document

ROOT = Path('/Users/mgolegario/bcc/itau_inovacamp/apf-hackathon')
OUT = ROOT / 'deliverables'
OUT.mkdir(exist_ok=True)


def set_text(paragraph, value: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = value
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(value)


def update(path: str, replacements: dict[int, str], output: Path, appendix=None) -> None:
    doc = Document(path)
    for index, value in replacements.items():
        set_text(doc.paragraphs[index], value)
    if appendix:
        fallback = doc.paragraphs[0].style
        heading1 = next(
            (p.style for p in doc.paragraphs if p.style and p.style.style_id.lower().startswith('heading1')),
            fallback,
        )
        heading2 = next(
            (p.style for p in doc.paragraphs if p.style and p.style.style_id.lower().startswith('heading2')),
            heading1,
        )
        list_style = next(
            (p.style for p in doc.paragraphs if p.style and p.style.style_id.lower().startswith('list')),
            fallback,
        )
        doc.add_page_break()
        title = doc.add_paragraph(appendix['title'])
        title.style = heading1
        for label, body in appendix['sections']:
            subtitle = doc.add_paragraph(label)
            subtitle.style = heading2
            if isinstance(body, list):
                for item in body:
                    bullet = doc.add_paragraph(item)
                    bullet.style = list_style
            else:
                doc.add_paragraph(body)
    doc.save(output)


def replace_table_text(path: Path, replacements: dict[str, str]) -> None:
    doc = Document(path)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text in replacements:
                        set_text(paragraph, replacements[paragraph.text])
    doc.save(path)


update(
    '/Users/mgolegario/Downloads/proposta-inovacamp-v9.docx',
    {
        37: 'Protótipo: serviço único em Python (FastAPI) + SQLite + Groq. A arquitetura separa orquestração, transações, decisão, visão de documentos, proatividade e canais; o estado financeiro é persistido no banco. Em produção, esses limites podem virar serviços sem reescrever as regras de negócio. Não são agentes autônomos independentes: é um único motor com roteamento de intenção e chamadas de função por módulo.',
        40: 'Opção Arrojada — Telegram integrado por webhook com segredo, vínculo do chat por código de uso único e estado persistente. Quando há ação financeira, o bot cria um token de 15 minutos e abre o app por deep link; o Telegram nunca confirma a transação.',
        43: 'Motor textual: Groq API com openai/gpt-oss-120b e fallback openai/gpt-oss-20b. Leitura de boleto: qwen/qwen3.6-27b quando a visão está disponível; fallback de palco determinístico e explicitamente rotulado quando não está.',
        44: 'Gerenciador de contexto: conversas, rascunhos transacionais, última transferência, boletos fotografados e tokens de continuação persistidos em SQLite.',
        49: 'Módulo de Pagamento: acessa saldo, fatura e limite sintéticos; compara Pix, débito e crédito mostrando impacto e premissas, sem inventar economia.',
        67: 'Otimização: APF compara Pix, débito e crédito com rich cards; mostra impacto em saldo, fatura e limite, além das premissas. Valores vêm das APIs ou de dados sintéticos rotulados no protótipo.',
        68: 'Execução: confirmação proporcional ao risco e ao canal — autenticação nativa do dispositivo dentro do app; fallback de demonstração identificado quando o ambiente não oferece biometria; deep link seguro quando a jornada começa no Telegram.',
        88: 'Opção Arrojada — Telegram integrado: vínculo por código temporário, webhook autenticado, estado persistente e handoff por token de uso único. Ações financeiras sempre continuam no app autenticado e nunca são concluídas no chat.',
        89: 'Evolução para WhatsApp: o adaptador de canal pode ser trocado sem alterar a orquestração. A homologação comercial e as políticas do provedor devem ser tratadas antes de produção.',
        93: 'Fadiga de notificação: cap de 2 notificações proativas por dia, janela de envio (9h–20h), cooldown por categoria (3 a 30 dias), consentimento granular e silenciamento por categoria. Agrupamento e ajuste por dismiss permanecem como evolução futura.',
        104: 'Hipótese de adoção a validar: até 480 mil clientes usando o comparador no primeiro ano. O protótipo não atribui economia de juros sem dados reais de comportamento e preço.',
        108: 'Premissas ilustrativas — base elegível, adoção, CAC, churn, ticket, spread e payback devem ser substituídos por dados internos do Itaú antes de qualquer decisão de investimento.',
        115: 'Toda ação que movimenta dinheiro exige confirmação no app autenticado. O consentimento LGPD é granular por categoria. O protótipo usa dados sintéticos rotulados, captura boleto por imagem, proatividade automática, persistência entre canais e integração Telegram pronta para receber credenciais. A arquitetura modular permite trocar os componentes de demonstração por serviços Itaú sem alterar a jornada.',
    },
    OUT / 'proposta-inovacamp-v10.docx',
    {
        'title': '12. Estado de Implementação e Evidências',
        'sections': [
            ('Implementado e testado', [
                'Pix/TED multi-turno, boleto, investimento, recibo e atualização de saldo.',
                'Captura de boleto por câmera/galeria, revisão dos campos e pagamento após confirmação.',
                'Biometria nativa do dispositivo, com fallback de demonstração claramente identificado.',
                'Telegram com vínculo, webhook autenticado, deep link, token de uso único e persistência de estado.',
                'Comparador com impacto em saldo, fatura e limite; premissas visíveis.',
                'Proatividade automática com consentimento, cap, janela, cooldown e mute.',
            ]),
            ('Hipóteses que ainda exigem validação', [
                'Aumento de engajamento, redução de atraso e diminuição de churn.',
                'Taxas de adoção, CAC evitado, spread, ticket, investimento e payback.',
                'Qualidade do scanner em um conjunto de boletos anonimizados e homologados.',
                'Atestação de dispositivo, antifraude e integração com APIs Itaú de produção.',
            ]),
        ],
    },
)

update(
    '/Users/mgolegario/Downloads/plano-arquitetura-apf-1dia-v2.docx',
    {
        3: 'Arquitetura mínima viável para ser construída em 1 dia por 2 pessoas. Eliminamos toda complexidade operacional desnecessária: um único serviço, um único banco de dados (SQLite), dados sintéticos explicitamente rotulados como tal e inferência via Groq, sujeita às cotas e condições vigentes do provedor.',
        7: 'LLM: Groq API — openai/gpt-oss-120b para linguagem, fallback 20b; qwen/qwen3.6-27b para visão de boleto quando disponível.',
        8: 'Interfaces: Flutter com câmera, autenticação do dispositivo e deep links; Telegram com webhook autenticado e vínculo por código temporário.',
        12: '''apf-hackathon/
├── main.py                      # FastAPI: chat, ações, scanner e canais
├── transfers.py                 # Pix/TED/boleto + estado transacional
├── document_intelligence.py     # Visão Groq + fallback de demo rotulado
├── telegram_integration.py      # Vínculo, webhook, handoff e tokens
├── models.py / database.py      # SQLite: dados, fluxos e identidades
├── decision.py / proactive.py   # Comparador, suitability e antecipação
├── llm.py / epc.py              # Linguagem com grounding + padrões SQL
├── configure_telegram.py        # Configuração do webhook
├── .env.example                 # Variáveis sem segredos
└── flutter_app/                 # Chat, câmera, biometria e app links''',
        15: 'O FastAPI expõe chat WebSocket e APIs para dashboard, padrões, privacidade, proatividade, scanner de boleto, vínculo Telegram, webhook e continuação segura.',
        16: '''@app.websocket('/ws')
@app.post('/user/{id}/boleto/scan')
@app.post('/user/{id}/proactive/scan')
@app.post('/user/{id}/telegram/link-code')
@app.post('/telegram/webhook')
@app.get('/continue/{token}')

# Regra central: canais externos iniciam a jornada;
# confirmação financeira acontece somente no app autenticado.''',
        20: '''// Toda resposta identifica a origem dos dados:
{
  "saldo": {"disponivel": 4523.87},
  "boletos": [{"id": "SABESP_001", "valor": 187.40}],
  "_source": "sintético"
}

// Scanner: "groq_vision" ou "demo_fallback" sempre visível ao usuário.''',
        38: '8. Fadiga de Notificação — Cap, Janela, Cooldown e Silenciamento',
        41: 'Janela de envio: 9h às 20h. Fora dela, o disparo é bloqueado e registrado; um agendador persistente é evolução de produção.',
        42: 'Agrupamento inteligente de múltiplos padrões no mesmo dia: evolução futura, não demonstrada pelo protótipo.',
        45: 'Adaptação automática pela taxa de dismiss: evolução futura; o protótipo já oferece silenciamento explícito por categoria.',
        48: 'ChatScreen usa WebSocket, cards acionáveis e scanner por câmera/galeria. local_auth solicita autenticação real em dispositivo compatível; web e ambientes sem suporte exibem fallback de demonstração rotulado. app_links retoma tokens do Telegram. A proatividade é verificada automaticamente após a conexão. Onboarding mantém três consentimentos LGPD granulares.',
        49: '10. Telegram — Estrutura Pronta para Integrar',
        50: 'Preencher TELEGRAM_BOT_TOKEN, TELEGRAM_BOT_USERNAME, TELEGRAM_WEBHOOK_SECRET e PUBLIC_BASE_URL no .env.',
        51: 'Executar python configure_telegram.py para registrar o webhook HTTPS.',
        52: 'Gerar POST /user/{id}/telegram/link-code; o cliente abre o link t.me e vincula o chat com código de uso único.',
        53: 'Ações financeiras geram botão “Continuar com segurança no app”, token de 15 minutos e deep link. O Telegram nunca executa dinheiro.',
        56: 'A arquitetura continua deliberadamente simples, mas agora cobre a jornada completa: estado persistente, scanner multimodal, autenticação do dispositivo e handoff Telegram→app. Componentes de demonstração são substituíveis por serviços Itaú sem alterar contratos ou regras de negócio.',
        61: 'Fadiga: cap diário, janela, cooldown por categoria e silenciamento em 1 clique estão implementados; agrupamento e adaptação por dismiss são evoluções futuras.',
        62: 'Custo de infraestrutura local próximo de zero, sujeito às cotas e condições dos provedores externos. Escalabilidade futura: trocar SQLite por banco gerenciado e adicionar filas/observabilidade sem reescrever a lógica de jornada.',
    },
    OUT / 'plano-arquitetura-apf-1dia-v3.docx',
    {
        'title': '13. Critérios de Aceite da Integração',
        'sections': [
            ('Antes da demonstração', [
                'Backend saudável e base sintética carregada.',
                'Build Flutter validada; câmera e autenticação conferidas no dispositivo.',
                'Bot vinculado ao usuário demo e webhook HTTPS configurado.',
                'Boleto sintético “não pagável” disponível na galeria.',
            ]),
            ('Antes de produção', [
                'Desativar TELEGRAM_ALLOW_DEMO_USER e exigir vínculo verificado.',
                'Usar cofre de segredos, atestação de dispositivo, idempotência distribuída e auditoria.',
                'Substituir dados sintéticos e execução local por APIs e sandbox Itaú.',
            ]),
        ],
    },
)

replace_table_text(
    OUT / 'plano-arquitetura-apf-1dia-v3.docx',
    {
        'Groq API (Llama 3.1 70B)': 'Groq API (gpt-oss-120b; fallback 20b)',
        'Free-tier: 20 req/min. Sem GPU, sem cartão, latência < 300ms.': 'Inferência hospedada; cotas, preço e latência dependem do plano vigente.',
        'Webhook gratuito; ngrok expõe localhost em 1 comando.': 'Webhook HTTPS; túnel local apenas para demonstração.',
        'Rich cards no Flutter: comparador de pagamento, biometria simulada, onboarding LGPD.': 'Rich cards, comparador, autenticação nativa com fallback de demo e onboarding LGPD.',
        'Telegram bot com ngrok. Deep link para ações financeiras. Teste end-to-end.': 'Telegram por webhook HTTPS, vínculo, deep link e teste end-to-end.',
    },
)
